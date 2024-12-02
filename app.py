
import streamlit as st
from jamaibase import JamAI, protocol as p
from docx import Document
from io import BytesIO
import random
import string
from PyPDF2 import PdfReader

jamai = JamAI(api_key="jamai_sk_cad07015753b6467de4c41c4186576f04834859ccf1effed", project_id="proj_2d329533d712840b679e9ccf")

# Function to extract text from PDF
def extract_text_from_pdf(pdf_file):
    pdf = PdfReader(pdf_file)
    text = ""
    for page in pdf.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

# Function to generate a random filename
def generate_random_filename(extension=".docx"):
    random_str = ''.join(random.choices(string.ascii_letters + string.digits, k=10))
    return f"final_report_{random_str}{extension}"

# Set up the Streamlit app
st.set_page_config(page_title="Talent Matchmaker", page_icon="😉")
st.title("HR's Matchmaker: Your next Talent is with us")
# Custom CSS to style the UI
st.markdown(
    """
    <style>
    /* Global styles */
    body {
        color: #f0f0f0;
        font-family: 'Arial', sans-serif;
    }

    h1, h2, h3, h4 {
        font-family: 'Helvetica', sans-serif;
        font-size: 28px;
    }

    /* Header Section */
    h1 {
        text-align: center;
        color: #FFA500;
        animation: fadeIn 2s ease-out;
    }

    /* Text animation */
    @keyframes fadeIn {
        0% { opacity: 0; transform: translateY(-20px); }
        100% { opacity: 1; transform: translateY(0); }
    }

    /* Button Style */
    .stButton>button {
        background-color: #28a745; /* Green color */
        color: white;
        border: none;
        border-radius: 5px;
        padding: 12px 25px;
        font-size: 16px;
        cursor: pointer;
        transition: all 0.3s ease;
    }

    .stButton>button:hover {
        background-color: #218838; /* Darker green on hover */
        transform: scale(1.05);
    }

    .stButton>button:active {
        background-color: #1e7e34; /* Even darker green when clicked */
        transform: scale(1.02);
    }

    /* Result Box Style */
    .generated-output {
        background-color: #444;
        padding: 20px;
        border-radius: 15px;
        margin-top: 30px;
        box-shadow: 0px 4px 20px rgba(0, 0, 0, 0.6);
        color: #f0f0f0;
        animation: fadeIn 2s ease-out;
    }

    /* Title within Result Box */
    .generated-output h4 {
        color: #FFA500;
        margin-bottom: 10px;
    }

    /* Subsection for better readability */
    .generated-output p {
        margin: 5px 0;
        font-size: 14px;
        line-height: 1.5;
    }

    /* File Upload Section */
    .stFileUploader {
        border: 2px solid #4CAF50;
        border-radius: 10px;
        padding: 10px;
        transition: border-color 0.3s ease;
    }

    .stFileUploader:hover {
        border-color: #45a049;
    }

    /* Custom Scrollbar for long content */
    ::-webkit-scrollbar {
        width: 8px;
    }

    ::-webkit-scrollbar-thumb {
        background-color: #4CAF50;
        border-radius: 10px;
    }

    ::-webkit-scrollbar-thumb:hover {
        background-color: #45a049;
    }

    /* Footer Button */
    .stButton>button {
        background-color: #007BFF;
        color: white;
        border-radius: 8px;
        padding: 12px 20px;
    }

    .stButton>button:hover {
        background-color: #0056b3;
    }

    </style>
    """,
    unsafe_allow_html=True
)

# Containers for inputs
with st.container():
    st.header("📄 Upload their CV and your Job Description. \n Trust us, we will do the work ;)")
    # Upload PDF CV
    cv_pdf = st.file_uploader("Upload CV (PDF format)", type="pdf")
    # Job Description input
    jobScope = st.text_area("Enter your Job Description")

# Action to process inputs
if st.button("Let's go!", use_container_width=True):
    if cv_pdf and jobScope:
        # Extract text from CV PDF
        cv_text = extract_text_from_pdf(cv_pdf)

        # Add rows to the existing table with the input data
        try:
            completion = jamai.add_table_rows(
                "action",
                p.RowAddRequest(
                    table_id="hrMatchmaker",
                    data=[{"cv": cv_text, "jobScope": jobScope}],
                    stream=False
                )
            )

            # Display the output generated in the columns
            if completion.rows:
                output_row = completion.rows[0].columns
                capable = output_row.get("capable")
                work_experience = output_row.get("workExp")
                rating = output_row.get("rating")
                matching_skills = output_row.get("validSkills")
                skills_not_matching = output_row.get("nullSkills")
                salary = output_row.get("salary")
                final_report = output_row.get("finalReport")

                st.subheader("Result:")
                st.markdown(
                    f"""
                    <div class="generated-output">
                        <h4>🙂‍↕️🙂‍↔️ Capability:</h4> <p>{capable.text if capable else 'N/A'}</p>
                        <h4>💼 Work Experience:</h4> <p>{work_experience.text if work_experience else 'N/A'}</p>
                        <h4>✨ Talent's Rating:</h4> <p>{rating.text if rating else 'N/A'}</p>
                        <h4>😀 Matching skills:</h4> <p>{matching_skills.text if matching_skills else 'N/A'}</p>
                        <h4>☹️ Lacking skills:</h4> <p>{skills_not_matching.text if skills_not_matching else 'N/A'}</p>
                        <h4>💰 Suggested Salary per month:</h4> <p>{salary.text if salary else 'N/A'}</p>
                        <h4>📝 Final Report:</h4> <p>{final_report.text if final_report else 'N/A'}</p>
                    </div>
                    """,
                    unsafe_allow_html=True
                )

                # Download the final report as a .docx file
                with st.container():
                    st.subheader("📩 Download a copy of the Final Report")
                    doc = Document()
                    doc.add_heading("Report", level=1)
                    # Final Report Section
                    doc.add_heading("Final Report", level=2)
                    doc.add_paragraph(final_report.text if final_report else 'N/A')
                    # Summary and Work Experience Section
                    doc.add_heading("Work Experience", level=2)
                    doc.add_paragraph(work_experience.text if work_experience else 'N/A')
                    # Skills Assessment Section
                    doc.add_heading("Skills Assessment", level=2)
                    doc.add_paragraph("Matching Skills:")
                    doc.add_paragraph(matching_skills.text if matching_skills else 'N/A')
                    doc.add_paragraph("Lacking skills:")
                    doc.add_paragraph(skills_not_matching.text if skills_not_matching else 'N/A')
                    doc.add_heading("Recommended Salary per month:")
                    doc.add_paragraph(salary.text if salary else 'N/A')

                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    st.download_button(
                        label="Get a copy 😺",
                        data=buffer,
                        file_name=generate_random_filename(),
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
            else:
                st.error("⚠️ Failed to get a response. Please try again.")
        except Exception as e:
            st.error(f"❌ An error occurred: {e}")
    else:
        st.warning("⚠️ Please upload a CV and enter a job description.")